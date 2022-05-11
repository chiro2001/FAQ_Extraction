#转换为纯文本

from docx import Document
import pandas as pd



import os

root = '/home/wangrui/2022项目/annoted/'

dirs = os.listdir(root)

# files = [os.listdir(d) for d in dirs if not d.endswith('.py')]
files = []
for d in dirs:
    if not d.endswith('.py') and not d.endswith('.json') and not d.endswith('.txt'):
        files = files+[root+'/'+d+'/'+f for f in os.listdir(d) if f.endswith('docx') ]
print(files)

for inpath in files:
    # inpath = '/home/wangrui/2022项目/2022-01-17 招证课题二前期资料/2022-01-04 样例1/A股新股/普通账户新股申购（A股）业务汇总.docx'

    outpath = '/'.join(inpath.split('/')[:-1])+'/'+inpath.split('/')[-1].replace('.docx','_pure_text.txt')
    print(outpath)
    doc = Document(inpath)

    # doc.styles["Heading 2"]
    # for p in doc.paragraphs:
    #     print(len(p.runs),p.runs[0].bold)
        
        
    #     s_name = p.style.name
    #     if s_name.startswith('Heading'):
    #         print(s_name, p.text,sep=':')

    #每一段的内容
    # for para in doc.paragraphs:
    #     print(para.text)

    # #每一段的编号、内容
    # for i in range(len(doc.paragraphs)):
    #     print(str(i),  doc.paragraphs[i].text)

    #表格
    # tbs = doc.tables
    # for tb in tbs:
    #     #行
    #     for row in tb.rows:    
    #         #列    
    #         for cell in row.cells:
    #             pass
    #             # print(cell.text)
    #             #也可以用下面方法
    #             '''text = ''
    #             for p in cell.paragraphs:
    #                 text += p.text
    #             print(text)'''


    pure_text = ''
    q_a = {}
    obj=doc
    f = False
    temp = ''
    for p in obj.paragraphs:
        runs = p.runs
        for run in runs:
            pure_text+=run.text
        pure_text+='\n'

        # if not f and len(runs)==1:
        #     f=True
        #     temp = runs[0].text
        #     continue
        # if f and len(runs)==1:
        #     f=False
        #     q_a[temp] = runs[0].text
        
        # if f and len(runs)==1:
        #     q_a[temp]+=runs[0].text
        #     continue
        # if f and len(runs) > 1:
        #     f=False
        # if not f and len(runs) > 1 :
        #     # f=True
        #     # temp = runs[0].text
        #     q_a[runs[0].text] =''.join([ runs[i].text for i in range(1, len(runs))]) 
        # # print(runs[0].text,len(runs))
    # qa = (pd.DataFrame({'question':q_a.keys(), 'answer':q_a.values()}))
    # qa.to_csv(outpath+'/ques_ans.csv',encoding='utf-8',index=None)
    
        
    with open(outpath,'w',encoding='utf-8') as f:
        f.write(pure_text)